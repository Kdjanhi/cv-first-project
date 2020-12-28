from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


Document = Document()

#profile picture
Document.add_picture('cvprofilepicture.png')

# name phone number and email details
name = input('What is your name?')
speak('Hello' + name + ' How are you today?')
phone_number = input('What is your phone number')
email = input('What is your email?')

Document.add_paragraph(
    name + '|' + phone_number + '|' + email )


# about me 
Document.add_heading('About me')
Document.add_paragraph(
    input('Tell me about yourself? ')
)

# Work experience
Document.add_heading('Work Experience')
p = Document.add_paragraph()

company = input('Enter company')
from_date = input('From Date')
to_date = input('End Date')

p.add_run(company + ' ').bold = True 
p.add_run(from_date + '-' + to_date + '\n').italics = True

experience_details = input(
        'Describe your experience at ')
p.add_run(experience_details)

# More Experiences
while True:
    has_more_experiences = input(
        'Do you have any more experiences? ')
    if has_more_experiences.lower() == 'yes':
        p = Document.add_paragraph()

        company = input('Enter company')
        from_date = input('From Date')
        to_date = input('End Date')

        p.add_run(company + ' ').bold = True 
        p.add_run(
            from_date + '-' + to_date + '\n').italics = True

        experience_details = input(
        'Describe your experience at ')
        p.add_run(experience_details)
    else:
        break

# Skills
Document.add_heading('Skills')
skills = input('Enter Skill')
p = Document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? ')
    if has_more_skills.lower() == 'yes':
        skills = input('Enter Skill ')
        p.Document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

    # footer
    section = Document.sections[0]
    footer = section.footer
    p = footer.paragraph [0]
    p.text = 'CV generated using VSC'
    




Document.save('cv.docx')